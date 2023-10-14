from docx import Document
from docx.oxml.ns import qn
import requests
from bs4 import BeautifulSoup
from docx.shared import RGBColor, Pt
import time
import os
import platform
import subprocess

def fetch_links_from_bing(query):
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }
    
    url = f"https://www.bing.com/search?q={query}"
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
    except requests.RequestException as e:
        print(f"Failed to fetch the search page. Error: {e}")
        return []
    
    soup = BeautifulSoup(response.content, 'html.parser')
    excluded_links = [
        "https://dxzhgl.miit.gov.cn/dxxzsp/xkz/xkzgl/resource/qiyereport.jsp?num=caf04fa4-bd8a-4d9e-80b6-2aa1b86c1509&type=yreport",
        "https://beian.miit.gov.cn",
        "https://go.microsoft.com/fwlink/?linkid=868922",
        "https://support.microsoft.com/topic/82d20721-2d6f-4012-a13d-d1910ccf203f&clcid=0x04"
    ]

    links = [a['href'] for a in soup.find_all('a', href=True) if a['href'].startswith('https') and a['href'] not in excluded_links]
    return links

def fetch_content_from_url(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
    except requests.RequestException as e:
        print(f"Failed to fetch the web page for {url}. Error: {e}")
        return None
    
    soup = BeautifulSoup(response.content, 'html.parser')
    return soup

def save_to_docx(soup, doc):
    # 设置正文字体为“宋体”
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(10.5)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    
    for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
        if tag.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag.name[1])
            head = doc.add_heading("", level=level)
            run = head.add_run(tag.text)
            run.font.name = u'Cambria'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            doc.add_paragraph(tag.text)

def open_file(filename):
    """打开文件的函数"""
    if platform.system() == "Windows":
        os.startfile(filename)
    elif platform.system() == "Darwin":  # macOS
        subprocess.Popen(["open", filename])
    elif platform.system() == "Linux":
        subprocess.Popen(["xdg-open", filename])
    else:
        print("Sorry, we don't support this operating system.")

def fetch_and_save_to_docx(query):
    filename = query.replace(" ", "_") + ".docx"
    
    links = fetch_links_from_bing(query)
    unique_links = set(links)  # 使用set来去除重复链接

    if not unique_links:
        print("No links fetched. Exiting...")
        return

    if os.path.exists(filename):
        doc = Document(filename)  # 加载已经存在的文件
    else:
        doc = Document()  # 创建一个新的Document对象

    for link in unique_links:
        print(f"Fetching content from: {link}")
        soup_content = fetch_content_from_url(link)
        
        if soup_content:
            save_to_docx(soup_content, doc)
        else:
            print(f"Skipping {link} due to fetch error.")

        time.sleep(2)  # 休眠3秒，避免请求太频繁

    doc.save(filename)
    print(f"File saved as {filename}")
    open_file(filename) 

if __name__ == '__main__':
    fetch_and_save_to_docx("创新")
