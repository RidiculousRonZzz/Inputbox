import openai
import os
import pyautogui
import time
import datetime
import json
import difflib
import tkinter as tk
import numpy as np
import pygetwindow as gw
from sklearn.metrics.pairwise import cosine_similarity
from tkinter import ttk, messagebox, Toplevel, Checkbutton
from pywinauto.application import Application
import time
import clipboard
from docx import Document
import pyperclip
import win32com.client
from everythingSearch import search_everything_results

from promptSave import *
from arxiv import *
from crawler import fetch_and_save_to_docx, open_file
from searchLastWeek import search_all_files

def get_embedding(text):
    api_key = "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"
    openai.api_key = api_key
    response = openai.Embedding.create(
        input=text,
        model="text-embedding-ada-002"
    )
    return np.array(response['data'][0]['embedding'])

def get_embeddings(texts):
    api_key = "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"
    openai.api_key = api_key
    responses = openai.Embedding.create(
        input=texts,
        model="text-embedding-ada-002"
    )
    return [np.array(response['embedding']) for response in responses['data']]

def gpt4_api(system, history):
    """ 返回str，参数为str,List """
    # api_key = os.getenv('OPENAI_API_KEY')
    api_key = "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"
    openai.api_key = api_key

    try:  # gpt-4-0314
        response = openai.ChatCompletion.create(model="gpt-4-0314", messages=[construct_system(system), *history])
        return response['choices'][0]['message']['content']
    except openai.error.ServiceUnavailableError:
        print('The server is overloaded or not ready yet. Please try again later.')
        return None
    except Exception as e:
        print(f'Unexpected error occurred: {e}')
        return None
    
def gpt4_api_1(system, history):
    """ 返回str，参数为str,List """
    api_key = "sk-1lfNSZtcTrbRTeSTW2GoT3BlbkFJMTL9f9PrpMULVIuwOzwz"
    openai.api_key = api_key

    try:
        response = openai.ChatCompletion.create(model="gpt-4", messages=[construct_system(system), *history])
        return response['choices'][0]['message']['content']
    except openai.error.ServiceUnavailableError:
        print('The server is overloaded or not ready yet. Please try again later.')
        return None
    except Exception as e:
        print(f'Unexpected error occurred: {e}')
        return None
    
def gpt4_api_2(system, history):
    """ 返回str，参数为str,List """
    api_key = "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"
    openai.api_key = api_key

    try:
        response = openai.ChatCompletion.create(model="gpt-4-0613", messages=[construct_system(system), *history])
        return response['choices'][0]['message']['content']
    except openai.error.ServiceUnavailableError:
        print('The server is overloaded or not ready yet. Please try again later.')
        return None
    except Exception as e:
        print(f'Unexpected error occurred: {e}')
        return None
    
def gpt4_api_3(system, history):
    """ 返回str，参数为str,List """
    # api_key = os.getenv('OPENAI_API_KEY')
    api_key = "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"
    openai.api_key = api_key

    try:  # gpt-4-0314
        response = openai.ChatCompletion.create(model="gpt-4-0314", messages=[construct_system(system), *history])
        return response['choices'][0]['message']['content']
    except openai.error.ServiceUnavailableError:
        print('The server is overloaded or not ready yet. Please try again later.')
        return None
    except Exception as e:
        print(f'Unexpected error occurred: {e}')
        return None
    
def gpt4_api_4(system, history):
    """ 返回str，参数为str,List """
    # api_key = os.getenv('OPENAI_API_KEY')
    api_key = "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"
    openai.api_key = api_key

    try:  # gpt-4-0314
        response = openai.ChatCompletion.create(model="gpt-4", messages=[construct_system(system), *history])
        return response['choices'][0]['message']['content']
    except openai.error.ServiceUnavailableError:
        print('The server is overloaded or not ready yet. Please try again later.')
        return None
    except Exception as e:
        print(f'Unexpected error occurred: {e}')
        return None
    
def turbo_api(system, history):
    """ 返回str，参数为str,List """
    api_key = "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"
    openai.api_key = api_key

    try:
        response = openai.ChatCompletion.create(model="gpt-3.5-turbo-16k-0613", messages=[construct_system(system), *history])
        return response['choices'][0]['message']['content']
    except openai.error.ServiceUnavailableError:
        print('The server is overloaded or not ready yet. Please try again later.')
        return None
    except Exception as e:
        print(f'Unexpected error occurred: {e}')
        return None

def construct_system(text):
    return construct_text("system", text)

def construct_text(role, text):
    return {"role": role, "content": text}

def construct_user(text):
    return construct_text("user", text)

def get_gpt4_recommendations(input_str):
    response = gpt4_api(PROMPT, [construct_user(input_str)])
    print(response)
    recommendations = {
        "应用": [],
        "网页": []
    }
    current_category = None

    for item in response.split("\n"):
        if item.strip().startswith("应用："):
            current_category = "应用"
        elif item.strip().startswith("网页："):
            current_category = "网页"
        else:
            if current_category and "：" not in item:
                clean_item = item.strip()
                if clean_item:
                    recommendations[current_category].append(clean_item)

    return recommendations

def create_outlook_reminder(subject, start, duration, reminder_minutes_before_start, body):
    # 获取Outlook应用程序的引用
    outlook = win32com.client.Dispatch("Outlook.Application")
    
    # 创建一个新的日历事件
    appointment = outlook.CreateItem(1)  # 1代表AppointmentItem

    # 设置事件的属性
    appointment.Subject = subject
    appointment.Start = start
    appointment.Duration = duration
    appointment.ReminderMinutesBeforeStart = reminder_minutes_before_start
    appointment.Body = body
    appointment.Save()

def open_app_if_exists(app_name):
    """模拟按下 Win+S，输入应用名称，再按回车来启动应用"""
    pyautogui.hotkey('win', 's')
    time.sleep(3)  # 等待搜索框出现
    pyperclip.copy(app_name)
    pyautogui.hotkey('ctrl', 'v')
    time.sleep(3)  # 给搜索时间来找到应用
    pyautogui.press('enter')
    pyautogui.press('enter')
    time.sleep(3)  # 等待应用启动或搜索完成

def search_user_input(filename, target_input):  # 模糊搜索
    close_matches = []
    old_input = ''
    with open(filename, 'r', encoding='utf-8') as file:
        all_user_inputs = [json.loads(line.strip())['user_input'] for line in file if line.strip()]
    matches = difflib.get_close_matches(target_input, all_user_inputs, n=1, cutoff=0.7)

    if matches:
        with open(filename, 'r', encoding='utf-8') as file:
            for line in file:
                try:
                    data = json.loads(line.strip())
                    if data['user_input'] in matches:
                        for decision in data['user_decision']:
                            close_matches.append(decision)
                            old_input = data['user_input']
                except json.JSONDecodeError:
                    continue

    return close_matches, old_input
    
# 从所有窗口中根据窗口句柄找到正确的窗口
def get_window_by_handle(handle):
    for window in gw.getWindowsWithTitle(''):
        if window._hWnd == handle:
            return window
    return None

def execute_application_or_web_search(category, recommendation, user_profile, user_input):
    """基于类别执行应用程序或在IE浏览器中搜索网页"""
    screen_width, screen_height = pyautogui.size()
    # 获取当前所有窗口的句柄
    initial_windows_handles = set(window._hWnd for window in gw.getWindowsWithTitle(''))

    if category == "应用":
        if recommendation.lower() in ["wechat", "weixin", "微信"]:
            wechat_message = gpt4_api(f"你是一位{user_profile}", [construct_user(f"你要{user_input}，打开微信，你会发什么")])
            print(wechat_message)
            send_wechat_message("ATM1", wechat_message)
        elif recommendation.lower() in ["excel", "excel.exe"]:
            excel_code = gpt4_api(EXCEL_PROMPT, [construct_user(f"输入：\n职业：{user_profile}\n任务：{user_input}\n应用：Excel")])
            print(excel_code)
            with open("excel_demo.py", "w", encoding="utf-8") as file:
                file.write(excel_code)
            os.system("python excel_demo.py")
        elif recommendation.lower() in ["word", "winword.exe"]:
            user_input_filename = user_input.replace(" ", "_") + ".docx"
            if os.path.exists(user_input_filename):
                doc = Document(user_input_filename)  # 加载已经存在的文件
            else:
                doc = Document()  # 创建一个新的Document对象open_file(filename) 
            doc.save(user_input_filename)
            open_file(user_input_filename)
        elif recommendation.lower() in ["日历", "outlook"]:
            OutlookReminderWindow(user_input)
        elif recommendation.lower() in ["arxiv"]:
            # 需要输入具体的名称，且是英文
            arxiv_fetch_and_save_to_excel_and_pdf(user_input)
        else:
            open_app_if_exists(recommendation)
            if not search_app_in_files(recommendation):
                search_in_ie_browser(recommendation)
        print(f"尝试打开应用：{recommendation}")
        new_windows_handles = set(window._hWnd for window in gw.getWindowsWithTitle('')) - initial_windows_handles
        for handle in new_windows_handles:
            window = get_window_by_handle(handle)
            if window:  # 确保找到了窗口
                window.moveTo(screen_width // 2, 0)
                window.resizeTo(screen_width // 2, screen_height)
    else:
        search_in_ie_browser(recommendation)
        judge_application_or_web = gpt4_api(APP_OR_WEB, [construct_user(recommendation)])
        if(judge_application_or_web == "0"):
            print(f"{recommendation}是文案素材")
            fetch_and_save_to_docx(recommendation)
        else:
            print(f"{recommendation}是应用")

def return_execute_application_or_web_search(category, recommendation, user_profile, user_input):
    """以前已经做过，现在只是恢复。基于类别执行应用程序或在IE浏览器中搜索网页"""
    screen_width, screen_height = pyautogui.size()
    # 获取当前所有窗口的句柄
    initial_windows_handles = set(window._hWnd for window in gw.getWindowsWithTitle(''))

    if category == "应用":
        if recommendation.lower() in ["wechat", "weixin", "微信"]:
            wechat_message = gpt4_api(f"你是一位{user_profile}", [construct_user(f"你要{user_input}，打开微信，你会发什么？请直接写出消息内容！！！")])
            print(wechat_message)
            send_wechat_message("ATM1", wechat_message)
        elif recommendation.lower() in ["excel", "excel.exe"]:
            os.system(f'start {user_input.replace(" ", "_")}.xlsx')
        elif recommendation.lower() in ["word", "winword.exe"]:
            os.system(f'start {user_input.replace(" ", "_")}.docx')
        elif recommendation.lower() in ["日历", "outlook"]:
            OutlookReminderWindow(user_input)
        else:
            open_app_if_exists(recommendation)
            if not search_app_in_files(recommendation):
                search_in_ie_browser(recommendation)
        print(f"尝试打开应用：{recommendation}")
        new_windows_handles = set(window._hWnd for window in gw.getWindowsWithTitle('')) - initial_windows_handles
        for handle in new_windows_handles:
            window = get_window_by_handle(handle)
            if window:  # 确保找到了窗口
                window.moveTo(screen_width // 2, 0)
                window.resizeTo(screen_width // 2, screen_height)
    else:
        judge_application_or_web = gpt4_api(APP_OR_WEB, [construct_user(recommendation)])
        if(judge_application_or_web == "0"):
            print(f"{recommendation}是文案素材")
            os.system(f'start {recommendation.replace(" ", "_")}.docx')
            time.sleep(3)
            print(f'打开相关文档：{recommendation.replace(" ", "_")}.docx')
        else:
            search_in_ie_browser(recommendation)
            print(f"{recommendation}是应用")

def search_in_ie_browser(query):
    """在IE浏览器中搜索给定的查询"""
    screen_width, screen_height = pyautogui.size()
    initial_windows_handles = set(window._hWnd for window in gw.getWindowsWithTitle(''))
    os.system(f'"C:\\Program Files\\Internet Explorer\\iexplore.exe" -new https://www.bing.com/search?q={query}')
    print(f"在IE浏览器中搜索：{query}")
    new_windows_handles = set(window._hWnd for window in gw.getWindowsWithTitle('')) - initial_windows_handles
    for handle in new_windows_handles:
        window = get_window_by_handle(handle)
        if window:  # 确保找到了窗口
            window.moveTo(0, 0)
            window.resizeTo(screen_width // 2, screen_height)

def save_user_choice(user_input, decisions_made):
    current_timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    data = {
        'timestamp': current_timestamp,
        'user_input': user_input,
        'user_decision': decisions_made
    }
    
    with open('user_choices.txt', 'a', encoding='utf-8') as file:
        file.write(json.dumps(data, ensure_ascii=False) + '\n')

def search_app_in_files(app_name):
    file_names = ["directory_structure.txt", "all_installed_apps.txt"]
    
    for file_name in file_names:
        with open(file_name, "r", encoding="utf-8", errors='replace') as f:
            contents = f.readlines()

        # 检查应用是否在文件中
        for line in contents:
            if app_name.strip().lower() in line.strip().lower():  # 使用lower()方法使搜索不区分大小写
                return True
    return False

def get_app_usage_count():
    app_counts = {}
    try:
        with open('user_choice.txt', 'r', encoding='utf-8') as file:
            for line in file:
                data = json.loads(line.strip())
                for decision in data['user_decision']:
                    if decision["category"] == "应用":
                        app_name = decision["recommendation"]
                        if app_name in app_counts:
                            app_counts[app_name] += 1
                        else:
                            app_counts[app_name] = 1
    except FileNotFoundError:
        print("'user_choice.txt' not found. Please ensure it's in the correct directory.")
    except json.JSONDecodeError:
        print("'user_choice.txt' contains invalid data.")
    
    return app_counts

def send_wechat_message(contact_name, message):
    """
    通过微信向指定的联系人发送消息。

    参数:
    - contact_name: 要搜索的联系人名称。
    - message: 要发送的消息。
    """
    # 启动微信（或者连接到正在运行的实例）
    app = Application(backend="uia").connect(path=r"D:\Program Files (x86)\Tencent\WeChat\WeChat.exe")

    # 等待微信界面加载
    dlg = app.window(title_re=".*Weixin.*")
    dlg.wait('visible', timeout=20)

    # 将微信窗口置于最上层
    dlg.set_focus()

    # 定位搜索框并输入联系人名称
    search_box = dlg.window(title="Search", control_type="Edit")
    search_box.click_input()
    time.sleep(3)
    search_box.type_keys(contact_name)
    time.sleep(3)

    # 定位到联系人的搜索结果
    target_contact = dlg.child_window(title=contact_name, control_type="ListItem", found_index=0)
    target_contact.wait('visible', timeout=10)

    # 点击联系人的搜索结果以进入聊天窗口，老是点不进去
    if target_contact.exists():
        target_contact.click_input()
        time.sleep(3)
    else:
        print("Target contact not found!")

    # 定位并发送消息
    edit_box = dlg.child_window(title=contact_name, control_type="Edit")
    if edit_box.exists():
        print("Found the edit box!")
        edit_box.click_input()
        time.sleep(3)
        
        lines = message.split('\n')
        for idx, line in enumerate(lines):
            clipboard.copy(line)
            edit_box.type_keys("^v")
            # 如果不是最后一行，就模拟按下 Shift + Enter 来换行
            if idx != len(lines) - 1:
                edit_box.type_keys("{VK_SHIFT down}{ENTER}{VK_SHIFT up}")
                time.sleep(1)
        
        # 按下 Enter 键来发送整个消息
        edit_box.type_keys("{ENTER}")  
    else:
        print("Failed to find the edit box.")

def update_usage_count(decisions_made):
    usage_file = "usage_counts.txt"
    if os.path.exists(usage_file):
        with open(usage_file, "r") as f:
            lines = f.readlines()
            usage_counts = {line.split(":")[0].strip(): int(line.split(":")[1].strip()) for line in lines}
    else:
        usage_counts = {}

    for decision in decisions_made:
        key = decision["category"] + "-" + decision["recommendation"]
        if key in usage_counts:
            usage_counts[key] += 1
        else:
            usage_counts[key] = 1

    # 对usage_counts进行排序
    sorted_usage_counts = sorted(usage_counts.items(), key=lambda x: (x[0].split('-')[0] != "应用", -x[1]))

    with open(usage_file, "w") as f:
        for key, count in sorted_usage_counts:
            f.write(f"{key}: {count}\n")

def split_string_into_segments(file_name, lines_per_segment=200):
    with open(file_name, 'r', encoding='utf-8') as file:
        lines = file.readlines()
    segments = []
    for i in range(0, len(lines), lines_per_segment):
        segment = ''.join(lines[i:i+lines_per_segment])
        segments.append(segment)
        print(f"segment:\n{segment}")
    return segments

class OutlookReminderWindow(Toplevel):
    def __init__(self, user_input):
        super().__init__()
        self.center_window()
        self.user_input = user_input
        self.title("设置提醒")
        self.geometry("450x600")
        
        self.subject_label = tk.Label(self, text="提醒主题")
        self.subject_label.pack(pady=5)
        self.subject_entry = tk.Entry(self)
        self.subject_entry.insert(0, user_input)
        self.subject_entry.pack(pady=5)

        self.start_label = tk.Label(self, text="事件开始时间（YYYY-MM-DD HH:MM）")
        self.start_label.pack(pady=5)
        self.start_entry = tk.Entry(self)
        self.start_entry.insert(0, "2023-10-06 14:00")
        self.start_entry.pack(pady=5)
        
        self.duration_label = tk.Label(self, text="事件时长（分钟）")
        self.duration_label.pack(pady=5)
        self.duration_entry = tk.Entry(self)
        self.duration_entry.insert(0, "60")
        self.duration_entry.pack(pady=5)
        
        self.reminder_label = tk.Label(self, text="提前几分钟提醒")
        self.reminder_label.pack(pady=5)
        self.reminder_entry = tk.Entry(self)
        self.reminder_entry.insert(0, "15")
        self.reminder_entry.pack(pady=5)
        
        self.body_label = tk.Label(self, text="提醒内容")
        self.body_label.pack(pady=5)
        self.body_entry = tk.Entry(self)
        self.body_entry.insert(0, user_input)
        self.body_entry.pack(pady=5)
        
        self.confirm_button = tk.Button(self, text="确认", command=self.confirm)
        self.confirm_button.pack(pady=20)
        
    def center_window(self, width=None, height=None):
        # 获取屏幕尺寸以计算布局属性中的位置
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        width = 450
        height = 600
        # 计算居中的坐标
        x = (screen_width - width) / 2
        y = (screen_height - height) / 2
        self.geometry('%dx%d+%d+%d' % (width, height, x, y))

    def confirm(self):
        start_time = self.start_entry.get()
        duration = int(self.duration_entry.get())
        reminder_minutes = int(self.reminder_entry.get())
        
        create_outlook_reminder(
            subject=self.user_input,
            start=start_time,
            duration=duration,
            reminder_minutes_before_start=reminder_minutes,
            body=self.user_input
        )
        self.destroy()
        open_app_if_exists("outlook")

class RecommendationsWindow(Toplevel):
    def __init__(self, recommendations, user_profile, user_input, parent):
        super().__init__(parent)
        self.title("推荐结果")
        self.center_window()
        self.recommendations = recommendations
        self.user_profile = user_profile
        self.user_input = user_input
        self.checked_items = []
        self.widgets = []

        self.canvas = tk.Canvas(self)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.scroll_frame = ttk.Frame(self.canvas)
        self.scroll_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.create_window((0, 0), window=self.scroll_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")

        for category, items in recommendations.items():
            cat_label = tk.Label(self.scroll_frame, text=category, font=("Arial", 12, "bold"))
            cat_label.pack(pady=12)
            self.widgets.append(cat_label)
            for recommendation in items:
                item_frame = ttk.Frame(self.scroll_frame)  # 为每个推荐项创建一个新的Frame
                item_frame.pack(fill='x', pady=5)
                
                var = tk.BooleanVar()
                checkbutton = Checkbutton(item_frame, variable=var)  # 父窗口改为item_frame
                checkbutton.pack(side="left", anchor='w')
                self.widgets.append(checkbutton)
                
                entry = tk.Entry(item_frame, width=30)  # 父窗口改为item_frame
                entry.insert(0, recommendation)
                entry.pack(side="left", anchor='w', padx=5)
                self.widgets.append(entry)
                
                self.checked_items.append((category, entry, var))

        # 手动添加推荐部分
        self.add_label = tk.Label(self.scroll_frame, text="手动添加推荐：")
        self.add_label.pack(pady=10)

        self.add_entry = tk.Entry(self.scroll_frame, width=30)
        self.add_entry.pack(pady=10)

        # 下拉菜单
        self.types = ["应用", "网页", "其他"]  # 推荐类型
        self.type_var = tk.StringVar()
        self.type_combobox = ttk.Combobox(self.scroll_frame, textvariable=self.type_var, values=self.types)
        self.type_combobox.set("应用")
        self.type_combobox.pack(pady=10)

        self.add_button = tk.Button(self.scroll_frame, text="添加", command=self.manual_add_recommendation)
        self.add_button.pack(pady=10)

        self.confirm_button = tk.Button(self.scroll_frame, text="确认", command=self.confirm_selections)  # 更改父窗口为 self.scroll_frame
        self.confirm_button.pack(pady=20)

    def manual_add_recommendation(self):
        new_recommendation = self.add_entry.get()
        category = self.type_var.get()
        if new_recommendation:
            var = tk.BooleanVar(value=True)  # 默认为选中状态
            checkbutton = Checkbutton(self.scroll_frame, text=new_recommendation, variable=var)
            checkbutton.pack(anchor='w')
            self.checked_items.append((category, new_recommendation, var))
            self.add_entry.delete(0, tk.END)

    def center_window(self, width=None, height=None):
        # 获取屏幕尺寸以计算布局属性中的位置
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        width = 450
        height = 900
        # 计算居中的坐标
        x = (screen_width - width) / 2
        y = (screen_height - height) / 2
        self.geometry('%dx%d+%d+%d' % (width, height, x, y))

    def confirm_selections(self):
        decisions_made = []
        for category, entry, var in self.checked_items:
            if var.get():
                # 检查entry是否是tk.Entry对象
                if isinstance(entry, tk.Entry):
                    recommendation = entry.get()  # 从tk.Entry控件中获取文本
                else:
                    recommendation = entry  # entry已经是一个字符串
                decisions_made.append({"category": category, "recommendation": recommendation})
                execute_application_or_web_search(category, recommendation, self.user_profile, self.user_input)

        save_decision = messagebox.askyesno("保存选择", "是否保存此次选择?")
        if save_decision:
            user_input = self.master.query_entry.get()
            save_user_choice(user_input, decisions_made)
            update_usage_count(decisions_made)

        # 关闭窗口
        self.destroy()

class AppGUI(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("万能输入框")
        self.center_window()
        # 添加标签和输入框 for 查询
        self.selection_var = tk.StringVar()
        self.selection_var.set("execute")  # 默认设置为执行任务

        # 使用Frame部件来容纳两个Radiobuttons
        self.radio_frame = tk.Frame(self)
        self.radio_frame.pack(padx=20, pady=5)

        # 在Frame部件中添加并排的Radiobutton
        self.execute_rb = tk.Radiobutton(self.radio_frame, text="执行任务", variable=self.selection_var, value="execute")
        self.execute_rb.grid(row=0, column=0, padx=10)  # 使用grid方法并指定row和column

        self.search_rb = tk.Radiobutton(self.radio_frame, text="搜索已有文件", variable=self.selection_var, value="search")
        self.search_rb.grid(row=0, column=1, padx=10)
        
        self.query_entry = tk.Entry(self, width=40)
        self.query_entry.pack(padx=20, pady=5)

        # 添加标签和输入框 for user_profile
        self.profile_label = tk.Label(self, text="请输入您的工作：")
        self.profile_label.pack(padx=20, pady=5)
        
        self.profile_entry = tk.Entry(self, width=40)
        self.profile_entry.insert(0, "学生")  # 默认值
        self.profile_entry.pack(padx=20, pady=5)
        
        # 添加查询按钮
        self.button = tk.Button(self, text="查询", command=self.on_query)
        self.button.pack(padx=20, pady=20)
    
    def center_window(self, width=None, height=None):
        # 获取屏幕尺寸以计算布局属性中的位置
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        width = 500
        height = 300
        # 计算居中的坐标
        x = (screen_width - width) / 2
        y = (screen_height - height) / 2
        self.geometry('%dx%d+%d+%d' % (width, height, x, y))

    def on_query(self):
        selected_option = self.selection_var.get()
        user_input = self.query_entry.get()
        user_profile = self.profile_entry.get()

        if not user_input or not user_profile:
                messagebox.showerror("错误", "请填写所有输入框")
                return
        
        if selected_option == "execute":      
            decisions_from_file, old_input = search_user_input('user_choices.txt', user_input)
            if decisions_from_file:
                print(f"'{old_input}' exists in 'user_choices.txt'")
                for decision in decisions_from_file:
                    category = decision["category"]
                    recommendation = decision["recommendation"]
                    return_execute_application_or_web_search(category, recommendation, user_profile, old_input)
                update_usage_count(decisions_from_file)
            else:
                recommendations = get_gpt4_recommendations(f"用户是一名{user_profile}。输入：{user_input}")
                RecommendationsWindow(recommendations, user_profile, user_input, self)
        elif selected_option == "search":
            print("Searching for existing files...")
            user_date = gpt4_api(IF_TIME, [construct_user(f"输入：{user_input}\n输出：\n")])
            print(user_date)
            if 'False' in user_date:  # 不涉及对时间的描述
                key_word = gpt4_api(KEY_WORD_SPLIT, [construct_user(f"输入：{user_input}\n输出：\n")])
                print(key_word)
                key_extension = gpt4_api(FILENAME_EXTENSION, [construct_user(f"输入：{user_input}\n输出：\n")])
                print(key_extension)
                words = key_word.strip().split("\n")
                big_dict = {}
                for word in words:
                    big_dict.update(search_everything_results(word, key_extension))
                keys_str = "\n".join(big_dict.keys())
                print(f"keys_str: {keys_str}")
                targetFile_output = gpt4_api(f"请找出“{user_input}”可能对应的文件，如果没有输出“无”。注意PPT也可能是.pdf格式", [construct_user(f"输入：{keys_str}\n输出：\n")])
                # 假设gpt4_api的输出是一个换行分隔的路径列表，我们将其拆分为单独的文件路径
                print(f"可能的文件：{targetFile_output}")
                individual_files = targetFile_output.split('\n')
                targetFiles_notime = []
                
                for targetFile in individual_files:
                    targetFile = targetFile.strip()  # 移除任何多余的空白字符，如换行符
                    if targetFile != "无":
                        targetFile = big_dict[targetFile]
                        if os.path.exists(targetFile):
                            subprocess.Popen(['start', '', targetFile], shell=True)
                            targetFiles_notime.append(targetFile)
                            print(f"{targetFile} exists!")
                        else:
                            print(f"{targetFile} does not exist!")
                    else:
                        print("无")
            else:  # 涉及对时间的描述
                start = ""
                end = ""
                search_code = gpt4_api(FILE_TIME_START_END, [construct_user(f"输入：{user_date}\n输出：\n")])
                print(search_code)
                exec(search_code)
                key_word = gpt4_api(KEY_WORD_SPLIT, [construct_user(f"输入：{user_input}\n输出：\n")])
                print(key_word)
                key_extension = gpt4_api(FILENAME_EXTENSION, [construct_user(f"输入：{user_input}\n输出：\n")])
                print(key_extension)
                words = key_word.strip().split("\n")
                big_dict = {}
                for word in words:
                    big_dict.update(search_everything_results(word, key_extension, start, end))
                keys_str = "\n".join(big_dict.keys())
                print(f"keys_str: {keys_str}")
                targetFile_output = gpt4_api(f"请找出“{user_input}”可能对应的文件，如果没有输出“无”。注意PPT也可能是.pdf格式", [construct_user(f"输入：{keys_str}\n输出：\n")])
                # 假设gpt4_api的输出是一个换行分隔的路径列表，我们将其拆分为单独的文件路径
                print(f"可能的文件：{targetFile_output}")
                individual_files = targetFile_output.split('\n')
                targetFiles = []
                
                for targetFile in individual_files:
                    targetFile = targetFile.strip()  # 移除任何多余的空白字符，如换行符
                    if targetFile != "无":
                        targetFile = big_dict[targetFile]
                        if os.path.exists(targetFile):
                            subprocess.Popen(['start', '', targetFile], shell=True)
                            targetFiles.append(targetFile)
                            print(f"{targetFile} exists!")
                        else:
                            print(f"{targetFile} does not exist!")
                    else:
                        print("无")




                # with open("searchLastWeek.py", "r", encoding="utf-8") as file:
                #     original_code = file.read()
                # with open("searchLastWeek.py", "a", encoding="utf-8") as file:
                #     file.write("\n"+search_code)
                #     file.write(OTHER_CODE)
                # os.system("python searchLastWeek.py")
                # # 恢复searchLastWeek.py到原始状态
                # with open("searchLastWeek.py", "w", encoding="utf-8") as file:
                #     file.write(original_code)
                # print("searchLastWeek.py 已恢复")

                # segments = split_string_into_segments("fileString.txt")
                        
                # for index, segment in enumerate(segments):
                #     print("可能有的")
                    
                #     if index % 6 == 0:
                #         targetFile_output = gpt4_api("请找出可能对应的文件，如果没有输出“无”\n"+segment, [construct_user(f"输入：{user_input}\n输出：\n")])
                #     elif index % 6 == 1:
                #         targetFile_output = gpt4_api_1("请找出可能对应的文件，如果没有输出“无”\n"+segment, [construct_user(f"输入：{user_input}\n输出：\n")])
                #     elif index % 6 == 2:
                #         targetFile_output = gpt4_api_2("请找出可能对应的文件，如果没有输出“无”\n"+segment, [construct_user(f"输入：{user_input}\n输出：\n")])
                #     elif index % 6 == 3:
                #         targetFile_output = gpt4_api_3("请找出可能对应的文件，如果没有输出“无”\n"+segment, [construct_user(f"输入：{user_input}\n输出：\n")])
                #     elif index % 6 == 4:
                #         targetFile_output = gpt4_api_4("请找出可能对应的文件，如果没有输出“无”\n"+segment, [construct_user(f"输入：{user_input}\n输出：\n")])
                #     else:
                #         targetFile_output = gpt4_api("请找出可能对应的文件，如果没有输出“无”\n"+segment, [construct_user(f"输入：{user_input}\n输出：\n")])
                    
                #     # 假设gpt4_api的输出是一个换行分隔的路径列表，我们将其拆分为单独的文件路径
                #     individual_files = targetFile_output.split('\n')
                    
                #     for targetFile in individual_files:
                #         targetFile = targetFile.strip()  # 移除任何多余的空白字符，如换行符
                #         if targetFile != "无":
                #             if os.path.exists(targetFile):
                #                 subprocess.Popen(['start', '', targetFile], shell=True)
                #                 targetFiles.append(targetFile)
                #                 print(f"{targetFile} exists!")
                #             else:
                #                 print(f"{targetFile} does not exist!")
                #         else:
                #             print("无")
                print(user_input + " " + "\n".join(targetFiles))

if __name__ == "__main__":
    app = AppGUI()
    app.mainloop()

# def get_answer_from_document(question, file_path):
#     doc = Document(file_path)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     document = '\n'.join(full_text)
#     # 使用OpenAI的embedding接口获取文档的向量表示
#     document_embedding = openai.Embedding.create(model="text-embedding-ada-002", input=document)

#     # 使用OpenAI的embedding接口获取问题的向量表示
#     question_embedding = openai.Embedding.create(model="text-embedding-ada-002", input=question)

#     # 计算问题和文档之间的相似度（这只是一个简单的示例，实际应用中可能需要更复杂的计算）
#     similarity_score = compute_similarity(document_embedding, question_embedding)
#     if similarity_score > 0.5:
#         response = gpt4_api(f"请根据这个文档回答问题：{document}", [construct_user(question)])
#         return response.choices[0].text.strip()
#     else:
#         return "Sorry, I couldn't find a relevant answer in the document."

# def compute_similarity(embedding1, embedding2):
#     # 这是一个简单的点积计算来比较两个向量的相似度
#     return sum([a*b for a, b in zip(embedding1, embedding2)])

# file_path = "《小丑》角色分析.docx"
# question = "文档是关于什么的？"
# print(get_answer_from_document(question, file_path))